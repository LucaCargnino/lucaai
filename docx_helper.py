import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report_from_markdown(markdown_text, output_path="Report_Operazioni.docx"):
    doc = Document()
    
    # Imposta margini standard (1 pollice / 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configura lo stile normale predefinito
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(40, 40, 40) # Grigio scuro per il testo normale

    lines = markdown_text.split('\n')
    in_table = False
    table_data = []

    def flush_table():
        nonlocal in_table, table_data
        if not table_data:
            return
        
        # Filtra le righe divisorie (tipo |---|---|)
        filtered_data = []
        for row in table_data:
            if all(re.match(r'^[\s:-|]+$', cell) for cell in row):
                continue
            filtered_data.append(row)
        
        if filtered_data:
            rows_count = len(filtered_data)
            cols_count = max(len(row) for row in filtered_data)
            
            table = doc.add_table(rows=rows_count, cols=cols_count)
            # Utilizza uno stile predefinito per le tabelle di Word
            table.style = 'Light Shading Accent 1'
            
            for r_idx, row in enumerate(filtered_data):
                for c_idx, cell in enumerate(row):
                    if c_idx < len(table.columns):
                        text = cell.strip()
                        # Se è la riga di intestazione, formattala in grassetto
                        cell_obj = table.cell(r_idx, c_idx)
                        cell_obj.text = text
                        if r_idx == 0:
                            for paragraph in cell_obj.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255) # Testo bianco per l'header della tabella
            
            # Aggiunge spazio vuoto dopo la tabella
            doc.add_paragraph()
            
        table_data = []
        in_table = False

    for line in lines:
        line_strip = line.strip()
        
        # Rilevamento righe di tabella markdown
        if line_strip.startswith('|') and line_strip.endswith('|'):
            in_table = True
            parts = line_strip.split('|')[1:-1]
            table_data.append(parts)
            continue
        elif in_table:
            flush_table()

        # Rilevamento Intestazioni
        if line_strip.startswith('# '):
            heading = doc.add_paragraph()
            run = heading.add_run(line_strip[2:])
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(31, 78, 121) # Blu scuro
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(8)
        elif line_strip.startswith('## '):
            heading = doc.add_paragraph()
            run = heading.add_run(line_strip[3:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(31, 78, 121)
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(6)
        elif line_strip.startswith('### '):
            heading = doc.add_paragraph()
            run = heading.add_run(line_strip[4:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(70, 70, 70)
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(4)
        
        # Elenchi puntati
        elif line_strip.startswith('- ') or line_strip.startswith('* '):
            content = line_strip[2:]
            p = doc.add_paragraph(style='List Bullet')
            # Cerca testo in grassetto
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)
            p.paragraph_format.space_after = Pt(4)
            
        # Linee vuote
        elif not line_strip:
            continue
            
        # Testo normale
        else:
            p = doc.add_paragraph()
            # Cerca testo in grassetto (**testo**)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)
            p.paragraph_format.space_after = Pt(6)
            
    # Se il file finisce con una tabella attiva
    if in_table:
        flush_table()

    # Assicuriamoci che la cartella di destinazione esista
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
        
    doc.save(output_path)
    print(f"[+] File Word salvato: {output_path}")
    return output_path

if __name__ == "__main__":
    # Test rapido di generazione
    test_md = """# Report Giornaliero Operazioni
## Account: test@example.com

Ecco la tabella delle operazioni estratte oggi:

| Priorità | Mittente | Oggetto | Azione Richiesta |
|---|---|---|---|
| **ALTA** | Mario Rossi | Preventivo logistica | Preparare preventivo entro sera |
| **MEDIA** | Supporto Cloud | Rinnovo dominio | Verificare scadenza fattura |

## Note aggiuntive
- Ricordarsi di inviare il file prima delle **18:00**.
- Verificare i dettagli con il team di marketing.
"""
    create_report_from_markdown(test_md, "test_output.docx")
